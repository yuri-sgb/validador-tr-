import streamlit as st
import re
import pandas as pd
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document

st.set_page_config(layout="wide")

st.title("Reorganizador Inteligente de Descritivo Técnico")

# ==============================
# ESCOLHA DA FONTE OFICIAL
# ==============================

fonte = st.radio(
    "Selecione a fonte oficial da ordem:",
    ["HTML (Tabela do TR)", "Manual"]
)

html_file = st.file_uploader(
    "Enviar HTML completo (tabela + descritivo)",
    type=["html", "txt"]
)

ordem_oficial = {}
ordem_confirmada = False

# ==============================
# LER HTML UMA ÚNICA VEZ
# ==============================

html_content = None

if html_file:
    html_content = html_file.read().decode("utf-8", errors="ignore")

# ==============================
# EXTRAÇÃO DA ORDEM PELO HTML
# ==============================

def extrair_ordem_html(html):

    soup = BeautifulSoup(html, "html.parser")

    tabela = soup.find("table")

    ordem = {}

    if tabela:

        linhas = tabela.find_all("tr")

        for linha in linhas:

            colunas = linha.find_all("td")

            if len(colunas) >= 3:

                numero = colunas[0].get_text(strip=True)
                catmat = colunas[2].get_text(strip=True)

                if numero.isdigit() and catmat.isdigit():

                    ordem[catmat] = int(numero)

    return ordem


# ==============================
# ORDEM MANUAL
# ==============================

if fonte == "Manual":

    texto_manual = st.text_area(
        "Cole no formato:\n1 - 430290\n2 - 250363",
        height=200
    )

    if st.button("Gerar Ordem Manual"):

        linhas = texto_manual.split("\n")

        for linha in linhas:

            if "-" in linha:

                partes = linha.split("-")

                numero = partes[0].strip()
                catmat = partes[1].strip()

                if numero.isdigit() and catmat.isdigit():

                    ordem_oficial[catmat] = int(numero)

        ordem_confirmada = True


# ==============================
# ORDEM VIA HTML
# ==============================

if fonte == "HTML (Tabela do TR)" and html_content:

    ordem_oficial = extrair_ordem_html(html_content)

    if ordem_oficial:

        df_ordem = pd.DataFrame(
            [{"CATMAT": k, "Número Oficial": v} for k, v in ordem_oficial.items()]
        ).sort_values("Número Oficial")

        st.subheader("Ordem detectada na Tabela")

        st.dataframe(df_ordem)

        if st.button("Confirmar Ordem da Tabela"):

            ordem_confirmada = True


# ==============================
# PROCESSAMENTO DO DESCRITIVO
# ==============================

def separar_blocos(html):

    return re.split(r'(?=<p class="Item_Nivel2">)', html)


def extrair_catmat(bloco):

    match = re.search(r'CATMAT:</strong>\s*(\d+)', bloco)

    return match.group(1) if match else None


def substituir_numero(bloco, novo_numero):

    return re.sub(
        r'(<strong[^>]*>Item\s*)\d+(:</strong>)',
        rf'\g<1>{novo_numero}\2',
        bloco,
        count=1
    )


# ==============================
# APLICAR CORREÇÃO
# ==============================

if ordem_confirmada and html_content:

    st.success("Ordem oficial confirmada.")

    if st.button("Aplicar Correção ao Descritivo"):

        blocos = separar_blocos(html_content)

        blocos_por_catmat = {}

        for bloco in blocos:

            catmat = extrair_catmat(bloco)

            if catmat:

                blocos_por_catmat[catmat] = bloco

        novo_html = ""

        for catmat, numero_correto in sorted(ordem_oficial.items(), key=lambda x: x[1]):

            if catmat in blocos_por_catmat:

                bloco_original = blocos_por_catmat[catmat]

                bloco_corrigido = substituir_numero(
                    bloco_original,
                    numero_correto
                )

                novo_html += bloco_corrigido

        st.subheader("Pré-visualização do HTML corrigido")

        st.code(novo_html[:3000], language="html")

        # ==============================
        # DOWNLOAD HTML
        # ==============================

        st.download_button(
            "Baixar HTML Corrigido",
            novo_html,
            file_name="descritivo_corrigido.html",
            mime="text/html"
        )

        # ==============================
        # GERAR WORD
        # ==============================

        doc = Document()

        doc.add_paragraph("Descritivo Técnico Corrigido")

        doc.add_paragraph(novo_html)

        buffer_word = BytesIO()

        doc.save(buffer_word)

        st.download_button(
            "Baixar Word",
            buffer_word.getvalue(),
            file_name="descritivo_corrigido.docx"
        )

        # ==============================
        # GERAR EXCEL ORDEM
        # ==============================

        df_export = pd.DataFrame(
            [{"CATMAT": k, "Número Oficial": v} for k, v in ordem_oficial.items()]
        ).sort_values("Número Oficial")

        buffer_excel = BytesIO()

        df_export.to_excel(buffer_excel, index=False)

        st.download_button(
            "Baixar Excel da Ordem Oficial",
            buffer_excel.getvalue(),
            file_name="ordem_oficial.xlsx"
        )
